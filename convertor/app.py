import os
import uuid
from flask import Flask, request, send_file, render_template, jsonify
from werkzeug.utils import secure_filename
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from PyPDF2 import PdfReader
from PIL import Image
from cryptography.fernet import Fernet

# Flask app
app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["CONVERTED_FOLDER"] = "converted"

# Create folders
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)
os.makedirs(app.config["CONVERTED_FOLDER"], exist_ok=True)

# ---------------- ENCRYPTION SETUP ----------------
KEY_FILE = "secret.key"

if not os.path.exists(KEY_FILE):
    with open(KEY_FILE, "wb") as f:
        f.write(Fernet.generate_key())

with open(KEY_FILE, "rb") as f:
    ENCRYPTION_KEY = f.read()

fernet = Fernet(ENCRYPTION_KEY)

def encrypt_file(input_path, output_path):
    with open(input_path, "rb") as f:
        data = f.read()
    encrypted = fernet.encrypt(data)
    with open(output_path, "wb") as f:
        f.write(encrypted)

def decrypt_file(input_path, output_path):
    with open(input_path, "rb") as f:
        data = f.read()
    decrypted = fernet.decrypt(data)
    with open(output_path, "wb") as f:
        f.write(decrypted)

# ---------------------------------------------------



@app.route("/")
def home():
    return render_template("converter.html")


@app.route("/decrypt", methods=["POST"])
def decrypt_file_route():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    encrypted_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(encrypted_path)

    # Decrypt file
    decrypted_name = "decrypted_" + uuid.uuid4().hex + "_" + filename.replace(".secure", "")
    decrypted_path = os.path.join(app.config["CONVERTED_FOLDER"], decrypted_name)
    
    try:
        decrypt_file(encrypted_path, decrypted_path)
        return send_file(
            decrypted_path,
            as_attachment=True,
            mimetype="application/octet-stream",
            download_name=decrypted_name
        )
    except Exception as e:
        return f"Decryption failed: {str(e)}", 400


# ---------- PDF → DOCX ----------
@app.route("/convert/pdf-to-docx", methods=["POST"])
def pdf_to_docx():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    pdf_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(pdf_path)

    # Create temp docx
    docx_temp = os.path.join(app.config["CONVERTED_FOLDER"], "temp_output.docx")

    reader = PdfReader(pdf_path)
    doc = Document()

    for page in reader.pages:
        text = page.extract_text()
        if text:
            doc.add_paragraph(text)

    doc.save(docx_temp)

    # Encrypt final file
    unique_name = uuid.uuid4().hex + ".secure"
    encrypted_path = os.path.join(app.config["CONVERTED_FOLDER"], unique_name)
    encrypt_file(docx_temp, encrypted_path)

    return send_file(
        encrypted_path,
        as_attachment=True,
        mimetype="application/octet-stream",
        download_name=unique_name
    )




# ---------- Word → PDF ----------
@app.route("/convert/word-to-pdf", methods=["POST"])
def word_to_pdf():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    filename = secure_filename(file.filename)
    word_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(word_path)

    pdf_temp = os.path.join(app.config["CONVERTED_FOLDER"], "temp_output.pdf")

    doc = Document(word_path)
    c = canvas.Canvas(pdf_temp, pagesize=letter)

    y = 750
    for para in doc.paragraphs:
        c.drawString(72, y, para.text)
        y -= 20
        if y < 50:
            c.showPage()
            y = 750

    c.save()

    # Encrypt
    unique_name = uuid.uuid4().hex + ".secure"
    encrypted_path = os.path.join(app.config["CONVERTED_FOLDER"], unique_name)
    encrypt_file(pdf_temp, encrypted_path)

    return send_file(
        encrypted_path,
        as_attachment=True,
        mimetype="application/octet-stream",
        download_name=unique_name
    )




# ---------- IMAGES → PDF ----------
@app.route("/convert/images-to-pdf", methods=["POST"])
def images_to_pdf():
    if "files" not in request.files:
        return jsonify({"error": "No images uploaded"}), 400

    files = request.files.getlist("files")
    if len(files) == 0:
        return jsonify({"error": "No images found"}), 400

    image_objects = []
    for f in files:
        img = Image.open(f).convert("RGB")
        image_objects.append(img)

    pdf_temp = os.path.join(app.config["CONVERTED_FOLDER"], "temp_images.pdf")

    if len(image_objects) == 1:
        image_objects[0].save(pdf_temp, "PDF")
    else:
        image_objects[0].save(pdf_temp, save_all=True, append_images=image_objects[1:])

    # Encrypt
    unique_name = uuid.uuid4().hex + ".secure"
    encrypted_path = os.path.join(app.config["CONVERTED_FOLDER"], unique_name)
    encrypt_file(pdf_temp, encrypted_path)

    return send_file(
        encrypted_path,
        as_attachment=True,
        mimetype="application/octet-stream",
        download_name=unique_name
    )


# ---------------------------------------------------

if __name__ == "__main__":
    app.run(host="127.0.0.1", port=7000, debug=True)
